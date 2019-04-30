#!/usr/bin/python3
# -*- coding: utf-8 -*-
""" This python script is written for the NPL challenge for April 2019. The goal is to read in
	a CSV file with information on WAPs extracted from Cisco Prime - the file is WAPS.csv. For the bonus
	challenge - the goal was to output a table with certain rows highlighted. I used the python-docx
	for this which allows a Word document to be created and formatting to be added to certain elements. I used
	this to build a table and highlight the appropriate rows.
	
	Arguments:
		1) Filename - (optional) with a relative path to the CSV input. If not there then WAPS.csv is sought
	
	Outputs:
		1) WAPS.docx - table report with certain rows highlighted
		
"""

# import modules HERE
import sys											# this allows us to analyze the arguments	
import os											# this allows us to check on the file
from datetime import datetime						# useful for getting timing information and for some data translation from Excel files
from contextlib import contextmanager
import statistics									# for calculating some basic statistics
import docx											# for building a Word document

# additional information about the script
__filename__ = "analyzeWAPS.py"
__author__ = "Robert Hallinan"
__email__ = "rhallinan@netcraftsmen.com"

#
# version history
#


"""
	20190430 - Initial version
"""

@contextmanager
def open_file(path, mode):
	the_file = open(path, mode)
	yield the_file
	the_file.close()

def importFile(passedArgs):
	""" this script will determine which function to run when parsing the input file, import the data, and return a list of dictionaries
	"""

 	# assign variables
	fileInput=str(passedArgs)

	# Does the file exist?
	if not os.path.exists(fileInput):
		print("File name provided to convert does not exist. Closing now...")
		sys.exit()

	if fileInput[-4:].lower() == ".csv":
		print("File Input is: "+fileInput)
		return parseCSV(fileInput)
	

def parseCSV (fileInput):
	""" This function parses the CSV file and returns a list of dictionaries with the keys of each dictionary as the column header and the value specific to the row
	"""

	print("Network information will be parsed from the CSV input file....")

	#
	# define outputs
	#

	# make a list for the items in the file
	outputNetDev=[]

	# open the input file for reading
	with open_file(fileInput,'r') as netDevFile:
		netDevFileLines=netDevFile.readlines()

	# set a trigger for first line and second lines - true at start, set to False when encountered
	firstLine = True

	# declare some general info so it is accessible for multiple iterations of the for loop once initially modified
	colHeaderList=[]

	for netDev in netDevFileLines:

		# declare empty dictionary that we can add to for this item's information
		newItem={}
		
		# skip the first line
		if firstLine:
			firstLine = False
			# Read in the second line as the column headers
			colHeaderRead=netDev.split(",")
			for f in range(len(colHeaderRead)):
				colHeaderList.append((f,colHeaderRead[f].rstrip().replace(u'\ufeff','')))
			continue

		# still going means that this is one of the entries
		
		# get a list of the line since is CSV
		itemList=netDev.split(",")
		
		# check the line - if the length of fields is longer than the length of columns then there is a comma somewhere in the entry
		# user has to make sure that there is no comma
		if len(itemList) != len(colHeaderList):
			print("One or more items have a comma in their value string which makes this impossible to properly parse as a CSV.")
			sys.exit()

		# get the info on this item
		for pair in colHeaderList:
			newItem[pair[1]]=str(itemList[pair[0]]).rstrip().replace(u'\ufeff','')
	
		# assign the dictionary of the new item to the list
		outputNetDev.append(newItem)

	return outputNetDev
	
def writeTableToWord(fileName, listOfLists):
	
	# make the document object
	document = docx.Document()
	
	# add initial header
	document.add_heading('AP Client Counts - Feb 2019', 0)
	
	# create a table
	table = document.add_table(rows=1, cols=len(listOfLists[0]))
	table.style = 'Table Grid'
	
	# set the header row
	hdr_cells = table.rows[0].cells
	i = 0
	for item in listOfLists[0]:
		hdr_cells[i].text = str(listOfLists[0][i])
		i += 1
	
	# identify the rows to be red
	redRows = [ '0', '1', '>= 30' ]
	
	# add the rest of the rows - set to red based on the above characteristic for the first item in the row
	thisRowNum = 1
	for indivRow in listOfLists[1:]:
		# add the row and return the tuple of cells
		row_cells = table.add_row().cells
		
		# for each element in this row
		thisItem = 0
		for item in listOfLists[thisRowNum]:
			# set the cell to the value
			run = row_cells[thisItem].add_paragraph().add_run(str(listOfLists[thisRowNum][thisItem]))
			
			# decide if the content needs to be set to red
			if listOfLists[thisRowNum][0] in redRows:
				font = run.font
				font.color.rgb = docx.shared.RGBColor(0xff, 0x00, 0x00)
			
			# remove the first paragraph (to add with content needed to add a paragraph - to delete is required for spacing)
			p = row_cells[thisItem].paragraphs[0]._element
			p.getparent().remove(p)
			p._p = p._element = None			

			# increment
			thisItem += 1
		
		# move on to the next row number
		thisRowNum += 1
	
	# save the Word file
	document.save(fileName)
	

def main(system_arguments):

	# get a python list of dictionaries by parsing the CSV file - validate that there is even an argument there using try
	try:
		fileName = system_arguments[1]
	except:
		fileName = "WAPS.csv"
	try:
		networkInventory = importFile(fileName)
	except:
		print(fileName + " does not exist in this directory. Exiting...")
		sys.exit()
	
	# initially filter the list to delete any items that don't exist
	networkInventory = [ x for x in networkInventory if x['AP Name'] != "" ]
	# print(networkInventory)
	
	# figure out the devices that are registered
	regDevices = [ x for x in networkInventory if x['Operational Status'] == "Registered" ]

	# figure out how many devices are not registered
	unregDevices = [ x for x in networkInventory if x['Operational Status'] == "Not Registered" ]
	
	# get name and IP of the unregistered devices
	unregNameIP = [ (x['AP Name'],x['IP Address']) for x in unregDevices ]
	newList = [['AP Name','IP Address']]
	for item in unregNameIP:
		newList += [list(item)]
	unregNameIP = newList
	# print(unregNameIP)
	
	# get total number of APs
	totalAPs = len(networkInventory)
	
	# get total number of clients
	clientsPerAP = [ int(x['Client Count']) for x in networkInventory ]
	totalClients = sum(clientsPerAP)
	# print(str(sum(clientsPerAP)))
	
	# clients with specific counts
	tableInfo = []
	tableInfo += [['# Clients on an AP', '# APs (Out of ' + str(totalAPs) + ')', '# Clients (Out of ' + str(totalClients) + ')']]
	
	# add information on APs with 0 clients
	zeroCli = len([ x for x in clientsPerAP if x == 0 ])
	tableInfo += [['0', str(zeroCli) + ' (' + '%.0f%%' % ( zeroCli / totalAPs * 100 ) + ')', '0' ]]
	
	# add information on APs with 1 client
	totalCurAP = len([ x for x in clientsPerAP if x == 1 ])
	totalCurCli = sum([ x for x in clientsPerAP if x == 1 ])
	tableInfo += [['1', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with 2 clients
	totalCurAP = len([ x for x in clientsPerAP if x == 2 ])
	totalCurCli = sum([ x for x in clientsPerAP if x == 2 ])
	tableInfo += [['2', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with 3 clients
	totalCurAP = len([ x for x in clientsPerAP if x == 3 ])
	totalCurCli = sum([ x for x in clientsPerAP if x == 3 ])
	tableInfo += [['3', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with 4 through 10 clients
	totalCurAP = len([ x for x in clientsPerAP if x >=4 and x <= 10 ])
	totalCurCli = sum([ x for x in clientsPerAP if x >=4 and x <= 10 ])
	tableInfo += [['4 - 10', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with 11 through 20 clients
	totalCurAP = len([ x for x in clientsPerAP if x >=11 and x <= 20 ])
	totalCurCli = sum([ x for x in clientsPerAP if x >=11 and x <= 20 ])
	tableInfo += [['11 - 20', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with 21 through 29 clients
	totalCurAP = len([ x for x in clientsPerAP if x >=21 and x <= 29 ])
	totalCurCli = sum([ x for x in clientsPerAP if x >=21 and x <= 29 ])
	tableInfo += [['21 - 29', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
	
	# add information on APs with more than 30 clients
	totalCurAP = len([ x for x in clientsPerAP if x >=30 ])
	totalCurCli = sum([ x for x in clientsPerAP if x >=30 ])
	tableInfo += [['>= 30', str(totalCurAP) + ' (' + '%.0f%%' % ( totalCurAP / totalAPs * 100 ) + ')', str(totalCurCli) + ' (' + '%.0f%%' % ( totalCurCli / totalClients * 100 ) + ')' ]]
		
	# provide some information
	print("The number of devices that are registered is: " + str(len(regDevices)))
	print("The number of devices that aren't registered is: " + str(len(unregDevices)))
	
	# print the table
	writeTableToWord('WAPS.docx', tableInfo)
		
if __name__ == "__main__":

	# this gets run if the script is called by itself from the command line
	main(sys.argv)